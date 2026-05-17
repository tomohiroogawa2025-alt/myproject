from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_ROOT = Path("outputs/reports/08_nl_mixture_ratio_estimation")
latest_dir = sorted([p for p in REPORT_ROOT.iterdir() if p.is_dir()])[-1]
summary = json.loads((latest_dir / "summary.json").read_text(encoding="utf-8"))
scores = []
with (latest_dir / "mixture_ratio_cv_scores.csv").open(encoding="utf-8-sig", newline="") as f:
    scores = list(csv.DictReader(f))

FONT_NAME = "Meiryo"
PRESET = {
    "page_margin_in": 1.0,
    "content_width_in": 6.5,
    "title_size": 20,
    "body_size": 10.5,
    "heading1_size": 15,
    "heading2_size": 12.5,
    "heading_blue": "2E74B5",
    "heading_dark": "1F4D78",
    "ink": "0B2545",
    "muted": "555555",
    "table_header_fill": "E8EEF5",
    "table_alt_fill": "F7F9FC",
}


def pct(x) -> str:
    return f"{float(x) * 100:.1f}%"


def pp(x) -> str:
    return f"{float(x):.2f} pt"


def set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, size=9, bold=bold, color=color)
    set_cell_margins(cell)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(
        run,
        size=PRESET["heading1_size"] if level == 1 else PRESET["heading2_size"],
        bold=True,
        color=PRESET["heading_blue"] if level == 1 else PRESET["heading_dark"],
    )
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, PRESET["body_size"])
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        set_font(run, 10)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="0B2545")
        set_cell_shading(hdr[i], PRESET["table_header_fill"])
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if r_i % 2 == 1:
                set_cell_shading(cells[i], PRESET["table_alt_fill"])
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path: Path, caption: str, width=6.2):
    if not path.exists():
        add_body(doc, f"図が見つかりません: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_font(r, 8.5, False, PRESET["muted"])


def configure_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    margin = Inches(PRESET["page_margin_in"])
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(PRESET["body_size"])
    return doc


def rows_for_range(range_name: str, n=8):
    rows = [r for r in scores if r["range"] == range_name]
    rows = sorted(rows, key=lambda r: float(r["rmse_percent_point"]))[:n]
    return [
        [
            r["preprocess"],
            r["model"],
            pp(r["rmse_percent_point"]),
            pp(r["mae_percent_point"]),
            f"{float(r['r2']):.3f}",
            pp(r["bias_percent_point"]),
            pct(r["within_5pt_rate"]),
            pct(r["within_10pt_rate"]),
        ]
        for r in rows
    ]


def comparison_rows():
    rows = []
    for name, label in [("full", "全波長"), ("900-1700nm", "900-1700 nm")]:
        b = summary["best_by_range"][name]
        rows.append(
            [
                label,
                f"{b['preprocess']} / {b['model']}",
                pp(b["rmse_percent_point"]),
                pp(b["mae_percent_point"]),
                f"{float(b['r2']):.3f}",
                pct(b["within_5pt_rate"]),
                pct(b["within_10pt_rate"]),
            ]
        )
    return rows


def moisture_error_rows(range_name: str):
    path = Path(summary["outputs"]["range_error_by_moisture"][range_name])
    rows = []
    with path.open(encoding="utf-8-sig", newline="") as f:
        for r in csv.DictReader(f):
            rows.append(
                [
                    r["moisture_bin"],
                    str(int(float(r["n"]))),
                    f"{float(r['moisture_min']):.1f}-{float(r['moisture_max']):.1f}",
                    f"{float(r['mean_moisture']):.1f}",
                    pp(r["rmse_percent_point"]),
                    pp(r["mae_percent_point"]),
                    pp(r["bias_percent_point"]),
                    pct(r["within_10pt_rate"]),
                ]
            )
    return rows


def build_report(range_name: str, title_label: str, out_name: str):
    best = summary["best_by_range"][range_name]
    figs = {k: Path(v) for k, v in summary["figures"].items()}
    range_figs = {k: Path(v) for k, v in summary["range_figures"][range_name].items()}
    error_csv = Path(summary["outputs"]["range_error_by_ratio"][range_name])
    error_rows = []
    with error_csv.open(encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            error_rows.append(row)

    doc = configure_doc()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run(f"N材/L材混合パルプ比率推定 可能性評価（{title_label}）")
    set_font(r, PRESET["title_size"], True, PRESET["ink"])
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(8)
    r = sub.add_run(
        f"08 簡易報告 | 解析run: {summary['run_id']} | 疑似混合スペクトル {summary['n_synthetic_spectra']:,}本 | 含水率近傍ペア {summary['n_matched_pairs']:,}組"
    )
    set_font(r, 9.5, False, PRESET["muted"])

    add_heading(doc, "要旨", 1)
    add_body(
        doc,
        "本評価では、N材/L材が混ざって作られたパルプシートの平均的な混合比率を、近赤外スペクトルから回帰モデルで推定できるかを予備検証した。既知混合比の実パルプシートがまだ無いため、含水率が近いN材・L材の単独スペクトルを線形混合し、疑似混合スペクトルを作成して評価した。",
    )
    add_bullets(
        doc,
        [
            f"{title_label}のベスト条件は {best['preprocess']} / {best['model']}。",
            f"RMSEは{pp(best['rmse_percent_point'])}、MAEは{pp(best['mae_percent_point'])}、R²は{float(best['r2']):.3f}。",
            f"±5 percentage point以内は{pct(best['within_5pt_rate'])}、±10 point以内は{pct(best['within_10pt_rate'])}。",
            "分類モデルの確率を比率とみなすのではなく、N材比率を目的変数にした回帰問題として扱う方針が妥当。",
            "本結果は疑似混合スペクトルに基づくため、実混合パルプシートの散乱、密度、坪量、叩解、繊維配向の影響は今後の実試料で検証が必要。",
        ],
    )

    add_heading(doc, "方法論", 1)
    add_body(
        doc,
        "パルプシート状態で既に混ざっている前提では、カメラで空間分布を見るよりも、測定したバルクスペクトルから平均的なN/L混合比を推定するモデル構築が主目的になる。したがって、最終的には既知配合比で作成したパルプシートを教師データとし、N材比率またはL材比率を目的変数にした回帰モデルを作る。",
    )
    add_bullets(
        doc,
        [
            "今回の予備検証: 含水率が近いN材/L材ペアを作成し、rawスペクトルをN比率0-100%で線形混合。",
            f"含水率マッチング条件: 許容差{summary['moisture_tolerance']}以内、差が小さい上位{summary['n_matched_pairs']}ペアを使用。",
            f"ペアの含水率差: 平均{summary['mean_pair_moisture_abs_diff']:.3f}、90%点{summary['p90_pair_moisture_abs_diff']:.3f}。",
            "評価: ペア単位のGroupKFoldにより、同じN/Lペア由来の混合比が学習・検証にまたがらないようにした。",
            "比較モデル: PLS回帰、Ridge回帰を中心に、raw/SNV/SNV+1次微分/SNV+2次微分を比較。",
        ],
    )
    add_figure(doc, figs["mixture_design"], "図1. N/Lペアの含水率差と疑似混合比率の分布", width=6.2)
    add_figure(doc, figs["moisture_distribution"], "図2. 線形合成に使った疑似混合スペクトルの含水率域", width=6.2)
    add_body(
        doc,
        f"線形合成に使った疑似混合スペクトルの推定含水率は、{summary['mixed_moisture_min']:.1f}から{summary['mixed_moisture_max']:.1f}の範囲で、平均{summary['mixed_moisture_mean']:.1f}、中央値{summary['mixed_moisture_median']:.1f}であった。低含水率側のデータ数が多く、高含水率側、特に120以上の領域はデータ数が少ないため、誤差比較の解釈ではサンプル数の偏りも考慮する。",
    )
    add_figure(doc, figs["example_mixed_spectra"], "図3. 含水率が近いN/Lペアから作成した疑似混合スペクトル例", width=6.2)

    add_heading(doc, "全波長版と900-1700 nm版の比較", 1)
    add_table(
        doc,
        ["範囲", "ベスト条件", "RMSE", "MAE", "R²", "±5pt以内", "±10pt以内"],
        comparison_rows(),
        widths=[1.0, 1.45, 0.8, 0.8, 0.65, 0.8, 0.8],
    )
    add_body(
        doc,
        f"全波長は{summary['n_full_features']}点、900-1700 nm相当は{summary['n_900_1700_features']}点を使用した。今回の疑似混合データでは全波長版の方が誤差は小さいが、900-1700 nm版でもR²が0.9を超えており、短波長側だけでも比率情報は一定程度残っている。",
    )
    add_figure(doc, figs["model_comparison"], "図4. 条件別CV RMSE上位モデル", width=6.2)

    add_heading(doc, f"{title_label}の推定結果", 1)
    add_table(
        doc,
        ["前処理", "モデル", "RMSE", "MAE", "R²", "bias", "±5pt以内", "±10pt以内"],
        rows_for_range(range_name),
        widths=[1.0, 0.9, 0.7, 0.7, 0.55, 0.7, 0.75, 0.75],
    )
    add_figure(doc, range_figs["actual_vs_predicted"], f"図5. {title_label}: 実際のN材比率 vs 推定N材比率", width=5.8)
    add_figure(doc, range_figs["error_by_ratio"], f"図6. {title_label}: N材比率ごとの推定誤差とMAE", width=6.2)

    add_heading(doc, "含水率域ごとの誤差", 1)
    add_body(
        doc,
        "線形合成時の推定含水率を0-15、15-30、30-60、60-120、120以上に分け、各含水率域で比率推定誤差を比較した。これは含水率が比率推定に与える影響を確認するための確認であり、実試料では含水率を制御した外部検証が必要である。",
    )
    add_table(
        doc,
        ["含水率域", "n", "範囲", "平均", "RMSE", "MAE", "bias", "±10pt以内"],
        moisture_error_rows(range_name),
        widths=[0.75, 0.45, 1.0, 0.6, 0.7, 0.7, 0.7, 0.75],
    )
    add_figure(doc, range_figs["error_by_moisture"], f"図7. {title_label}: 含水率域ごとのRMSE/MAE", width=6.2)
    add_body(
        doc,
        "今回の疑似混合評価では、低含水率域ほど誤差が小さく、高含水率域ほど誤差が大きくなる傾向が見られた。特に120以上の領域ではサンプル数が少ないうえに誤差が大きく、比率推定モデルを実運用するには高含水率域の校正データを意図的に増やす必要がある。",
    )

    add_heading(doc, "比率ごとの誤差", 1)
    add_table(
        doc,
        ["N材比率", "平均推定", "MAE", "bias", "95%絶対誤差"],
        [
            [
                f"{float(r['n_ratio_percent']):.0f}%",
                f"{float(r['mean_pred']):.1f}%",
                pp(r["mae"]),
                pp(r["bias"]),
                pp(r["p95_abs_error"]),
            ]
            for r in error_rows
        ],
        widths=[0.9, 0.9, 0.75, 0.75, 1.0],
    )
    add_body(
        doc,
        "端点付近では予測値を0-100%にクリップしているため、誤差の出方が中央比率と異なる可能性がある。実試料では0/100、100/0だけでなく、10/90や90/10のような端点近傍を十分に含めて検量線を確認する必要がある。",
    )

    add_heading(doc, "解釈と次アクション", 1)
    add_bullets(
        doc,
        [
            "今回の結果は、N/Lのスペクトル差が比率情報として使える可能性を示す予備的な根拠になる。",
            "ただし、疑似混合は単独スペクトルの線形合成であり、実際の混合パルプシートに固有の散乱・密度・繊維配向・叩解条件はまだ入っていない。",
            "次段階では、既知N/L比率のパルプシートを複数ロット・複数含水率で作成し、同じ手順で外部検証する。",
            "含水率が比率推定を邪魔する可能性があるため、含水率を制御した系列と、含水率を意図的に振った系列の両方が望ましい。",
            "運用指標は、RMSE/MAEに加えて、±5 point以内または±10 point以内で合格とするかを現場許容差から決める。",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run(f"08 mixture ratio estimation | {title_label}")
    set_font(r, 8, False, "777777")

    out_path = latest_dir / out_name
    doc.save(out_path)
    return out_path


full_path = build_report("full", "全波長版", "08_NL混合比率推定_全波長版_上司報告.docx")
nm_path = build_report("900-1700nm", "900-1700nm版", "08_NL混合比率推定_900-1700nm版_上司報告.docx")
print(full_path)
print(nm_path)
