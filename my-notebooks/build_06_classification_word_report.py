from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_ROOT = Path("outputs/reports/06_classification_possibility")
latest_dir = sorted([p for p in REPORT_ROOT.iterdir() if p.is_dir()])[-1]
summary = json.loads((latest_dir / "summary.json").read_text(encoding="utf-8"))
figs = {k: Path(v) for k, v in summary["figure_paths"].items()}
OUT_DOCX = latest_dir / "06_分類可能性評価_上司報告.docx"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


model_scores = read_csv_rows(latest_dir / "classification_cv_scores.csv")
top_wavenumbers = read_csv_rows(latest_dir / "top_discriminative_wavenumbers.csv")
pairwise = read_csv_rows(latest_dir / "species_pairwise_separability.csv")
species_counts = read_csv_rows(latest_dir / "species_counts.csv")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Hiragino Sans"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color="1F4D78" if level > 1 else "2E74B5")
    return p


def add_body(doc, text, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        set_font(r1, 11, True)
        r2 = p.add_run(text[len(bold_label) :])
        set_font(r2, 11)
    else:
        run = p.add_run(text)
        set_font(run, 11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, 10.5)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_figure(doc, path: Path, caption: str, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, 9, False, "555555")
    cap.paragraph_format.space_after = Pt(8)


def pct(x):
    return f"{float(x) * 100:.1f}%"


def wavelength_nm(wavenumber):
    return 1e7 / float(wavenumber)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Hiragino Sans"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("近赤外スペクトルによる樹種・N/L材分類可能性評価")
set_font(run, 20, True, "0B2545")
subtitle = doc.add_paragraph()
r = subtitle.add_run(f"簡易報告 | 解析run: {summary['run_id']} | サンプル数 {summary['n_samples']:,}、波数点 {summary['n_spectral_features']:,}")
set_font(r, 10, False, "555555")

add_heading(doc, "要旨", 1)
add_body(
    doc,
    "近赤外スペクトルから樹種13分類およびN/L材2分類の可能性を確認した。代表的な前処理・分類器による3-fold CVでは、両タスクとも非常に高い分類性能を示した。",
)
add_bullets(
    doc,
    [
        f"樹種分類の代表CV最良: balanced accuracy {pct(summary['species_best']['balanced_accuracy'])}（{summary['species_best']['preprocess']} / {summary['species_best']['model']}）。",
        f"N/L材分類の代表CV最良: balanced accuracy {pct(summary['kl_best']['balanced_accuracy'])}（{summary['kl_best']['preprocess']} / {summary['kl_best']['model']}）。",
        "N/L差は、ANOVA F値および標準化平均差の上位波数として整理した。候補領域は木材中のセルロース・ヘミセルロース・リグニン由来のC-H/O-H結合振動に関係する可能性がある。",
        "平均差だけでなくクラス内分散を考慮した標準化平均差でも、同領域はN/L分類に寄与しやすい候補として確認された。",
        f"PCAではPC1+PC2で{(summary['pca_metrics']['pc1_explained_variance'] + summary['pca_metrics']['pc2_explained_variance']) * 100:.1f}%、PC1-5累積で{summary['pca_metrics']['pc5_cumulative_explained_variance'] * 100:.1f}%を説明し、N/LのPC5シルエットは{summary['pca_metrics']['kl_silhouette_pc5']:.2f}であった。",
        "ただし、同一データセット内のCVであり、測定日・ロット・個体差をまたいだ外部検証ではないため、実運用可能性は独立検証で確認する必要がある。",
    ]
)

add_heading(doc, "データと分類ラベル", 1)
add_body(
    doc,
    f"入力は train.csv のスペクトル列のみを特徴量として使用した。含水率列は特徴量に含めていない。波数範囲は {summary['axis_min']:.0f}-{summary['axis_max']:.0f} cm⁻¹（約{wavelength_nm(summary['axis_max']):.0f}-{wavelength_nm(summary['axis_min']):.0f} nm）で、列名が数値の1555点をスペクトル特徴量として扱った。",
)
kl_rows = [[item["kl"], str(item["n"])] for item in summary["kl_counts"]]
add_table(doc, ["分類", "サンプル数"], kl_rows, widths=[1.4, 1.2])
add_figure(doc, figs["counts"], "図1. 樹種別およびN/L材別サンプル数", width=6.2)

add_heading(doc, "分類可能性の概観", 1)
add_body(
    doc,
    "PCAでは、樹種およびN/L材で一定の分離傾向が見られる。特にN/L材は大分類であり、樹種分類よりも低次元空間で分かれやすい。一方、樹種間では近い領域に重なる組み合わせもあり、ペアごとの確認が必要である。PCAは教師なしの可視化であるため、モデルの過学習とは別に、スペクトルそのものに群構造があるかを説明しやすい。",
)
add_figure(doc, figs["pca"], "図2. SNV後スペクトルのPCAスコアプロット", width=6.3)
add_figure(doc, figs["pca_scree_loadings"], "図3. PCA寄与率とPC1/PC2ローディング", width=6.2)
add_figure(doc, figs["pca_kl_centroids"], "図4. PCA上のN/L材セントロイド", width=5.2)

pca_metric_rows = [
    ["PC1寄与率", f"{summary['pca_metrics']['pc1_explained_variance'] * 100:.1f}%"],
    ["PC2寄与率", f"{summary['pca_metrics']['pc2_explained_variance'] * 100:.1f}%"],
    ["PC1-5累積寄与率", f"{summary['pca_metrics']['pc5_cumulative_explained_variance'] * 100:.1f}%"],
    ["N/L silhouette PC1-2", f"{summary['pca_metrics']['kl_silhouette_pc2']:.2f}"],
    ["N/L silhouette PC1-5", f"{summary['pca_metrics']['kl_silhouette_pc5']:.2f}"],
    ["樹種 silhouette PC1-5", f"{summary['pca_metrics']['species_silhouette_pc5']:.2f}"],
]
add_table(doc, ["PCA指標", "値"], pca_metric_rows, widths=[2.4, 1.5])
add_body(
    doc,
    "PCAシルエットは1に近いほど低次元空間上で群が分かれていることを示す。N/LのPC1-5での値が高い場合、N/L差が単なる分類器依存ではなく、スペクトル空間上の構造としても見えていると説明できる。",
)
if "pc1_moisture_corr" in summary["pca_metrics"]:
    add_figure(doc, figs["pca_moisture"], "図5. PCAスコアと含水率の関係", width=6.1)
    moisture_rows = [
        [f"PC{i}", f"{summary['pca_metrics'].get(f'pc{i}_moisture_corr', float('nan')):.3f}"]
        for i in range(1, 6)
    ]
    add_table(doc, ["PCA成分", "含水率との相関"], moisture_rows, widths=[1.4, 1.7])
    add_body(
        doc,
        f"今回のデータではPC1と含水率の相関は{summary['pca_metrics']['pc1_moisture_corr']:.3f}で中程度にとどまり、PC2と含水率の相関が{summary['pca_metrics']['pc2_moisture_corr']:.3f}とより強い。したがって、PCA上の水分影響はPC1よりもPC2方向に強く表れている可能性が高い。樹種/N-L分類の解釈では、PC2方向の分離が含水率差を反映していないかを別途確認する必要がある。",
    )

add_figure(doc, figs["cv_scores"], "図6. 代表モデルによる3-fold CV分類性能", width=6.3)

top_score_rows = []
for task in ["species", "kl"]:
    rows = [r for r in model_scores if r["task"] == task]
    rows = sorted(rows, key=lambda r: float(r["balanced_accuracy"]), reverse=True)[:3]
    for r in rows:
        top_score_rows.append([task, r["preprocess"], r["model"], pct(r["balanced_accuracy"]), pct(r["f1_macro"])])
add_table(doc, ["タスク", "前処理", "モデル", "Balanced acc.", "Macro F1"], top_score_rows, widths=[0.8, 1.5, 1.4, 1.2, 1.0])

add_heading(doc, "スペクトル差と分散の解釈", 1)
add_body(
    doc,
    "N/L材別の平均スペクトルでは、rawでも前処理後でも群間差が確認できる。ただし平均差だけでは、同じクラス内のばらつきに対して差が十分大きいか判断できない。そのため、平均±1SDと標準化平均差（L-N平均差 / pooled SD）を併せて確認した。",
)
add_body(
    doc,
    "スペクトル図は波数の向きは従来表示のままとし、上軸に波長nmを併記した。2次微分スペクトルおよび分散図は、外れ値的に大きい端部の影響で全体が潰れないよう、表示範囲を分位点ベースで制限している。",
)
add_figure(doc, figs["kl_mean_spectra"], "図7. N/L材別平均スペクトル", width=6.3)
add_figure(doc, figs["kl_mean_sd"], "図8. N/L材別平均±1SDスペクトル", width=6.3)
add_figure(doc, figs["kl_difference"], "図9. L材-N材の平均差スペクトル", width=6.3)
add_figure(doc, figs["kl_standardized_diff"], "図10. 標準化平均差とN/L材内ばらつき", width=6.1)

std_rows = []
for r in summary["top_kl_standardized_differences"][:8]:
    w = float(r["wavenumber_cm-1"])
    std_rows.append([f"{w:.0f}", f"{wavelength_nm(w):.0f}", f"{float(r['standardized_l_minus_k']):.2f}", f"{float(r['pooled_std_snv_deriv2']):.4f}"])
add_table(doc, ["波数 cm⁻¹", "波長 nm", "標準化差", "pooled SD"], std_rows, widths=[1.0, 1.0, 1.1, 1.2])
add_body(
    doc,
    "標準化平均差が大きい波数は、平均差が大きいだけでなくクラス内分散に対しても差が安定している領域である。分類可能性を説明する際は、平均差スペクトルよりもこちらの方が根拠として強い。",
)

kl_top_rows = []
for r in [x for x in top_wavenumbers if x["task"] == "kl"][:8]:
    w = float(r["wavenumber_cm-1"])
    if 4350 <= w <= 4550:
        interp = "C-H/O-H結合帯。セルロース・リグニン構成差の候補"
    elif 5750 <= w <= 5900:
        interp = "C-H第1倍音近傍。リグニン/抽出成分差の候補"
    elif 5000 <= w <= 5300:
        interp = "O-H結合帯。水分影響も受けやすい"
    else:
        interp = "木材成分差・散乱差の候補"
    kl_top_rows.append([f"{w:.0f}", f"{wavelength_nm(w):.0f}", f"{float(r['f_value']):.1f}", interp])
add_table(doc, ["波数 cm⁻¹", "波長 nm", "F値", "解釈メモ"], kl_top_rows, widths=[0.9, 0.9, 0.9, 3.8])
add_figure(doc, figs["top_kl_wavenumbers"], "図11. N/L分類で差が大きい波数候補", width=5.8)

add_heading(doc, "混同と注意点", 1)
add_body(
    doc,
    "混同行列では代表CV上はほぼ完全分類となった。ただし、この結果は分類可能性を示す一方で、測定条件や個体差がCV分割内に共有されている場合、性能が楽観的になる可能性がある。",
)
add_figure(doc, figs["species_cm"], "図12. 樹種分類の混同行列", width=5.9)
add_figure(doc, figs["kl_cm"], "図13. N/L材分類の混同行列", width=4.5)

hard_rows = []
for r in summary["hard_species_pairs"][:6]:
    hard_rows.append([r["class_a"], r["class_b"], pct(r["balanced_accuracy"])])
add_table(doc, ["樹種A", "樹種B", "ペア分離 balanced acc."], hard_rows, widths=[2.0, 2.0, 1.7])
add_body(
    doc,
    "ペア評価ではベイスギを含む組み合わせで分離が相対的に低い結果が出た。全体モデルでは高精度でも、個別ペアの難易度を見ると、類似スペクトルまたは前処理・次元圧縮条件に敏感な組み合わせが存在する可能性がある。",
)

add_heading(doc, "N材/L材混合パルプ比率推定に向けた考え方", 1)
add_body(
    doc,
    "本件の最終目標が、N材とL材が混ざってできたパルプ中のN材/L材比率推定である場合、現在の分類可能性評価は第一段階の確認に位置づけられる。分類でN/Lを分けられることは必要条件に近いが、それだけでは混合比を定量できるとは限らない。",
)
add_bullets(
    doc,
    [
        "分類モデルの確率値は、N材が何%含まれるかを直接表すものではない。確率は分類器の確信度であり、物理的な混合比とは別物である。",
        "混合比推定には、既知のN/L配合比で作った標準サンプルが必要である。例: N=0, 10, 20, ..., 100% の段階サンプルを複数ロット・複数含水率で測定する。",
        "目的変数をN材比率またはL材比率としたPLS回帰、Ridge/SVR/GPR、RandomForest回帰などで定量モデルを作る。評価指標はRMSE、MAE、R2、許容誤差内率を使う。",
        "粉砕度、繊維長、密度、含水率、散乱状態がスペクトルに強く効くため、比率以外の要因を実験計画で振るか、前処理と外部検証で頑健性を確認する必要がある。",
        "単一点またはバルク平均スペクトルなら、得られるのは測定視野全体の平均的な混合比である。空間的な分布や面積比を出すには、ハイパースペクトル/マルチスペクトル画像でピクセルごとの推定が必要になる。",
    ],
)
add_body(
    doc,
    "したがって次段階の可能性検証では、N材100%、L材100%、既知混合比サンプルを用意し、分類ではなく比率回帰として評価するのが望ましい。分類可能性評価で見えた有効波長領域は、比率回帰の波長選択候補として再利用できる。",
)

add_heading(doc, "レビューコメントと次アクション", 1)
add_bullets(
    doc,
    [
        "分類可能性は高い。樹種分類・N/L材分類とも、スペクトルに分類情報が十分含まれている可能性が高い。",
        "一方で、現在の結果は同一データ内の交差検証であり、未知ロット・未知測定日・未知個体への一般化性能はまだ未確認。",
        "含水率は学習特徴量に入れていないが、スペクトルには水分由来吸収が含まれるため、含水率差が分類に寄与している可能性は別途確認が必要。",
        "次は GroupKFold（ロット・個体・測定日単位）または外部検証データで再評価し、N/L材分類が樹種を覚えているだけでないかを検証する。",
        "もし面内のN/L材割合を評価したい場合は、通常の一点スペクトルではなく、ハイパースペクトル/マルチスペクトル画像でピクセルごとに分類して面積比を出す設計が必要。",
    ]
)

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("06 classification possibility brief")
set_font(r, 8, False, "777777")

doc.save(OUT_DOCX)
print(OUT_DOCX)
